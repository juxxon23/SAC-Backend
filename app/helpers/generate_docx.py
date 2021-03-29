from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate

class GenerateDocx():

    document = None
    tpl = None

    def new_doc(self):
        self.document = Document()

    def create_head(self):
        header = self.document.sections[0].header
        paragraph = header.paragraphs[0]
        logo = paragraph.add_run()
        logo.add_picture('logosena.png', width=Inches(5.77))
        self.document.add_page_break()

    def create_foot(self):
        footer = self.document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.text = "GD-F-007\nV03"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def get_content_table(self):
        for i in self.content:
            if type(i) == str:
                print("i:", type(i), i, sep="**")
            elif type(i) == list:
                for j in i:
                    if type(j) == list:
                        for k in j:
                            if type(k) == str:
                                print("k:", type(k), k, sep="**")
                    elif type(j) == str:
                        print("j:", type(j), j, sep="**")

    def select_template(self, docx_path):
        self.tpl = DocxTemplate(docx_path)

    def save_doc(self, n):
        doc_name = 'acta-{}.docx'.format(n)
        self.document.save(doc_name)

    def construct(self):
        # self.insert_content()
        # tpl.render(self.table)
        # tpl.save('acta01.docx')
        self.new_doc()
        self.create_head()
        self.create_foot()
        # self.create_table()
        # self.create_body()
        self.save_doc('1')


#docx_tool = GenerateDocx()
#docx_tool.construct()
